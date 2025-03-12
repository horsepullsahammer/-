import os
import requests
import urllib3
import json
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import logging
from datetime import datetime
from tkinter import Tk, Label, Button, OptionMenu, StringVar, ttk
from tkinter.messagebox import showerror
from pathlib import Path
import threading

# 配置日志记录，设置日志级别为INFO，并定义日志格式
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# 禁用 HTTPS 证书验证警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# 定义API的URL，用于获取影院排片信息
api_url = "https://apis.netstart.cn/maoyan/cinema/shows?cinemaId=26501&ci=1&channelId=4"

# 获取当前用户的桌面路径
def get_desktop_path():
    """获取当前用户的桌面路径，返回字符串形式的路径"""
    return str(Path.home() / "Desktop")

# 发送GET请求并返回响应数据
def fetch_data(url):
    """发送HTTP GET请求到指定URL，并返回解析后的JSON数据"""
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
        "Accept": "application/json",
        "Authorization": "Bearer YOUR_API_KEY",  # 如果需要认证
    }
    try:
        response = requests.get(url, headers=headers, verify=False)  # 禁用SSL证书验证
        response.raise_for_status()  # 如果HTTP状态码不是200，抛出异常
        return response.json()  # 返回解析后的JSON数据
    except requests.exceptions.RequestException as e:
        logging.error(f"请求失败: {e}")  # 记录请求失败的错误日志
        return None
    except json.JSONDecodeError as e:
        logging.error(f"JSON解析失败: {e}")  # 记录JSON解析失败的错误日志
        return None

# 从API数据中提取可用的放映日期列表
def extract_available_dates(data):
    """从API返回的数据中提取所有可用的放映日期，并返回排序后的日期列表"""
    if not data or 'data' not in data or 'movies' not in data['data']:
        return []  # 如果数据无效或缺少必要字段，返回空列表
    
    dates = set()  # 使用集合去重
    for movie in data['data']['movies']:
        for show in movie.get('shows', []):  # 遍历每部电影的放映信息
            show_date = show.get('showDate')  # 获取放映日期
            if show_date:
                dates.add(show_date)  # 添加到日期集合中
    return sorted(dates)  # 返回排序后的日期列表

# 将日期字符串转换为中文格式
def get_chinese_date(date_str):
    """将日期字符串（如2023-03-09）转换为中文格式（如3月9日（周日））"""
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")  # 将字符串解析为日期对象
        chinese_month = date_obj.month  # 获取月份
        chinese_day = date_obj.day      # 获取日期
        chinese_weekdays = {
            "Monday": "周一", "Tuesday": "周二", "Wednesday": "周三",
            "Thursday": "周四", "Friday": "周五", "Saturday": "周六", "Sunday": "周日"
        }
        chinese_weekday = chinese_weekdays[date_obj.strftime("%A")]  # 获取中文星期几
        return f"{chinese_month}月{chinese_day}日（{chinese_weekday}）"  # 返回中文格式的日期
    except Exception as e:
        logging.error(f"日期转换失败: {e}")  # 记录日期转换失败的错误日志
        return date_str  # 如果转换失败，返回原始日期字符串

# 创建Word文档，包含指定日期的电影排片信息
def create_document(data, selected_date):
    """根据API数据创建Word文档，仅包含所选日期的电影信息"""
    if not data or 'data' not in data:
        logging.error("无效的数据结构")  # 记录错误日志
        return None

    # 创建Word文档对象
    document = Document()

    # 设置页边距
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.2)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(1.17)
        section.right_margin = Cm(1.17)

    # 设置全局字体为仿宋
    style = document.styles['Normal']
    font = style.font
    font.name = '仿宋'
    font.size = Pt(18)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 添加标题
    title = f"{get_chinese_date(selected_date)}"  # 使用中文日期
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 提取电影信息并添加到文档中
    movies = data['data'].get('movies', [])
    for movie in movies:
        # 提取电影的基本信息
        movie_name = movie.get('nm')
        movie_duration = movie.get('dur')
        movie_genres = movie.get('desc', '').split('|')[1] if len(movie.get('desc', '').split('|')) > 1 else "未知类型"

        # 检查当天是否有放映场次
        shows = movie.get('shows', [])
        has_shows = any(show.get('showDate') == selected_date for show in shows)

        if not has_shows:
            continue  # 如果当天没有放映时间，跳过该电影

        # 提取放映类型（3D、2D 或 3D/2D）
        show_types = set()
        for show in shows:
            if show.get('showDate') == selected_date:
                for p in show.get('plist', []):
                    tp = p.get('tp')
                    if tp:
                        show_types.add(tp)

        # 根据放映类型设置前缀
        type_prefix = "3D/2D" if "3D" in show_types and "2D" in show_types else "3D" if "3D" in show_types else "2D"

        # 添加电影信息到文档
        movie_info = document.add_paragraph()
        # 添加放映类型前缀
        type_prefix_run = movie_info.add_run(f"{type_prefix}")
        type_prefix_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        type_prefix_run.font.name = '仿宋'
        type_prefix_run.font.size = Pt(18)
        type_prefix_run.bold = True

        # 添加电影名称
        movie_name_run = movie_info.add_run(f"《{movie_name}》")
        movie_name_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        movie_name_run.bold = True

        # 添加电影类型和时长
        movie_details_run = movie_info.add_run(f" {movie_genres} {movie_duration}分钟")
        movie_details_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        movie_details_run.bold = True

        # 添加放映时间和影厅号
        formatted_times = []
        for show in shows:
            if show.get('showDate') == selected_date:
                plist = show.get('plist', [])
                for p in plist:
                    show_time = p.get('tm')  # 获取放映时间
                    hall = p.get('th')       # 获取影厅号
                    if show_time and hall:
                        hall_number = ''.join(filter(str.isdigit, hall))  # 提取影厅号中的数字
                        formatted_times.append(f"{show_time}（{hall_number}）")  # 格式化时间和影厅号

        if formatted_times:
            times_paragraph = document.add_paragraph()
            for i, time in enumerate(formatted_times):
                time_run = times_paragraph.add_run(time)
                time_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                time_run.bold = True
                if (i + 1) % 6 == 0:  # 每6个时间换行
                    times_paragraph = document.add_paragraph()

        # 在每部电影的信息后添加一个空行
        document.add_paragraph()  # 添加空段落

    # 设置段落行间距
    for paragraph in document.paragraphs[1:]:
        paragraph.paragraph_format.line_spacing = 0.9
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    return document

# 保存文档，如果文件名冲突则自动重命名
def save_document_with_unique_name(desktop_path, filename):
    """保存文档，如果文件名冲突则自动重命名"""
    file_path = Path(desktop_path) / filename
    counter = 1
    original_filename, file_extension = os.path.splitext(filename)

    while file_path.exists():
        new_filename = f"{original_filename}_{counter}{file_extension}"
        file_path = Path(desktop_path) / new_filename
        counter += 1

    return file_path

# 打开生成的文档
def open_document(file_path):
    """打开生成的文档"""
    try:
        os.startfile(file_path)
    except Exception as e:
        logging.error(f"无法打开文件: {e}")

# 主应用程序类
class App:
    def __init__(self, root):
        self.root = root
        self.root.title("中影影城(临清店)排片生成器")
        self.root.geometry("300x220")

        # 显示今天的日期
        today = datetime.now()
        chinese_weekdays = {
            "Monday": "周一", "Tuesday": "周二", "Wednesday": "周三",
            "Thursday": "周四", "Friday": "周五", "Saturday": "周六", "Sunday": "周日"
        }
        chinese_weekday = chinese_weekdays[today.strftime("%A")]  # 获取中文星期几
        today_str = f"今天是{today.month}月{today.day}日（{chinese_weekday}）"  # 格式化日期

        # 添加标签显示今天的日期
        today_label = Label(root, text=today_str, font=("仿宋", 12))
        today_label.pack(pady=10)

        # 初始化界面组件
        self.status_label = Label(root, text="正在加载数据...")
        self.status_label.pack(pady=10)

        self.selected_date = StringVar(root)
        self.date_menu = OptionMenu(root, self.selected_date, "")
        self.date_menu.pack(pady=10)

        # 添加重试按钮，初始状态为隐藏
        self.retry_button = Button(root, text="重试", command=self.load_data, state="disabled")
        self.retry_button.pack_forget()

        self.generate_and_open_button = Button(root, text="生成文档并打开", command=self.on_generate_and_open, state="disabled")
        self.generate_and_open_button.pack(pady=10)

        self.data = None
        self.available_dates = []

        # 启动后台线程加载数据
        threading.Thread(target=self.load_data, daemon=True).start()

    def load_data(self):
        """在后台加载数据"""
        self.status_label.config(text="正在加载数据，请稍候...")
        self.retry_button.config(state="disabled")
        self.generate_and_open_button.config(state="disabled")

        def fetch_data_in_thread():
            self.data = fetch_data(api_url)
            if not self.data:
                self.status_label.config(text="无法获取影城排期数据，请稍后再试！")
                self.root.after(0, self.show_retry_button)
                return

            self.available_dates = extract_available_dates(self.data)
            if not self.available_dates:
                self.status_label.config(text="API数据中没有可用的放映日期！")
                self.root.after(0, self.show_retry_button)
                return

            # 将日期转换为中文格式
            chinese_dates = [get_chinese_date(date) for date in self.available_dates]

            self.selected_date.set(self.available_dates[0])  # 设置默认日期为第一个日期
            self.date_menu['menu'].delete(0, 'end')  # 清空日期选择列表
            for date, chinese_date in zip(self.available_dates, chinese_dates):
                # 将中文日期显示在界面上，但实际保存的值为原始日期
                self.date_menu['menu'].add_command(label=chinese_date, command=lambda value=date: self.selected_date.set(value))

            self.status_label.config(text="数据加载完成")
            self.root.after(0, self.hide_retry_button)
            self.generate_and_open_button.config(state="normal")

        threading.Thread(target=fetch_data_in_thread, daemon=True).start()

    def show_retry_button(self):
        """显示重试按钮"""
        self.retry_button.pack(pady=5)
        self.retry_button.config(state="normal")

    def hide_retry_button(self):
        """隐藏重试按钮"""
        self.retry_button.pack_forget()

    def on_generate_and_open(self):
        """生成文档并打开"""
        selected_date_value = self.selected_date.get()
        document = create_document(self.data, selected_date_value)
        if not document:
            showerror("错误", "无法生成文档！")
            return

        output_filename = f"中影影城排片_{selected_date_value.replace('-', '')}.docx"
        saved_file_path = save_document_with_unique_name(get_desktop_path(), output_filename)
        document.save(saved_file_path)

        open_document(saved_file_path)

# 主函数
def main():
    root = Tk()
    app = App(root)

    # 获取屏幕的宽度和高度
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()

    # 计算窗口的初始位置
    window_width = 350  # 窗口的宽度
    window_height = 220  # 窗口的高度
    x_position = (screen_width // 2) - (window_width // 2)
    y_position = (screen_height // 2) - (window_height // 2)

    # 设置窗口的初始位置
    root.geometry(f"{window_width}x{window_height}+{x_position}+{y_position}")

    root.mainloop()

if __name__ == "__main__":
    main()